"""
Complete RAG System with Google Gemini LLM Integration
Includes: Document Upload → Chunking → Embedding → Storage → Query → LLM Response
Database: SQLite for metadata storage
"""

# ============================================================================
# IMPORTS
# ============================================================================

# Standard Library
import asyncio
import logging
import re
import tempfile
import shutil
import json
import uuid
import sqlite3
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Optional
from contextlib import contextmanager

# FastAPI
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

# Document Processing
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import pandas as pd

# Embeddings
from sentence_transformers import SentenceTransformer

# Vector Store
from pinecone import Pinecone, ServerlessSpec
import numpy as np

# Google Gemini
import google.generativeai as genai

# Environment Variables
from dotenv import load_dotenv
import os

# Load environment variables
load_dotenv()

# ============================================================================
# CONFIGURATION
# ============================================================================

# API Keys
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY", "")
PINECONE_INDEX_NAME = os.getenv("PINECONE_INDEX_NAME", "rag-documents")

# Constants
MAX_DOCUMENTS = int(os.getenv("MAX_DOCUMENTS", "20"))
MAX_PAGES_PER_DOC = int(os.getenv("MAX_PAGES_PER_DOC", "1000"))
MAX_FILE_SIZE_MB = int(os.getenv("MAX_FILE_SIZE_MB", "50"))
SUPPORTED_FORMATS = set(os.getenv("SUPPORTED_FORMATS", ".pdf,.docx,.txt,.html,.htm,.csv,.json,.md,.xml").split(","))

# Chunking and Embedding
DEFAULT_CHUNK_SIZE = int(os.getenv("DEFAULT_CHUNK_SIZE", "512"))
DEFAULT_CHUNK_OVERLAP = int(os.getenv("DEFAULT_CHUNK_OVERLAP", "50"))
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "all-MiniLM-L6-v2")

# Gemini Settings
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-1.5-flash")
MAX_CONTEXT_CHUNKS = int(os.getenv("MAX_CONTEXT_CHUNKS", "5"))

# Database
DATABASE_PATH = os.getenv("DATABASE_PATH", "data/metadata.db")

# Configure logging
logging.basicConfig(
    level=getattr(logging, os.getenv("LOG_LEVEL", "INFO")),
    format=os.getenv("LOG_FORMAT", '%(asctime)s - %(name)s - %(levelname)s - %(message)s')
)
logger = logging.getLogger(__name__)

# Initialize FastAPI
app = FastAPI(
    title="Complete RAG System with Gemini LLM",
    description="Upload documents, search with semantic similarity, and get AI-powered responses using Google Gemini",
    version="3.1.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=os.getenv("CORS_ALLOW_ORIGINS", "*").split(","),
    allow_credentials=os.getenv("CORS_ALLOW_CREDENTIALS", "true").lower() == "true",
    allow_methods=os.getenv("CORS_ALLOW_METHODS", "*").split(","),
    allow_headers=os.getenv("CORS_ALLOW_HEADERS", "*").split(","),
)


# ============================================================================
# DATABASE SETUP
# ============================================================================

class DatabaseManager:
    """Manages SQLite database for document metadata."""

    def __init__(self, db_path: str):
        self.db_path = db_path
        Path(db_path).parent.mkdir(parents=True, exist_ok=True)
        self._init_database()

    @contextmanager
    def get_connection(self):
        """Context manager for database connections."""
        conn = sqlite3.connect(self.db_path)
        conn.row_factory = sqlite3.Row
        try:
            yield conn
            conn.commit()
        except Exception as e:
            conn.rollback()
            raise e
        finally:
            conn.close()

    def _init_database(self):
        """Create tables if they don't exist."""
        with self.get_connection() as conn:
            cursor = conn.cursor()

            # Documents table
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS documents (
                    document_id TEXT PRIMARY KEY,
                    filename TEXT NOT NULL,
                    file_type TEXT NOT NULL,
                    file_size_mb REAL,
                    page_count INTEGER,
                    word_count INTEGER,
                    character_count INTEGER,
                    upload_timestamp TEXT,
                    status TEXT,
                    error_message TEXT
                )
            ''')

            # Chunks table
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS chunks (
                    chunk_id TEXT PRIMARY KEY,
                    document_id TEXT NOT NULL,
                    chunk_index INTEGER,
                    text TEXT,
                    word_count INTEGER,
                    embedding_generated BOOLEAN,
                    stored_in_vectordb BOOLEAN,
                    created_timestamp TEXT,
                    FOREIGN KEY (document_id) REFERENCES documents (document_id)
                )
            ''')

            # Queries table
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS queries (
                    query_id TEXT PRIMARY KEY,
                    query_text TEXT NOT NULL,
                    query_timestamp TEXT,
                    num_results INTEGER,
                    llm_response TEXT,
                    processing_time_ms REAL
                )
            ''')

            # Query results table
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS query_results (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    query_id TEXT NOT NULL,
                    chunk_id TEXT NOT NULL,
                    relevance_score REAL,
                    FOREIGN KEY (query_id) REFERENCES queries (query_id),
                    FOREIGN KEY (chunk_id) REFERENCES chunks (chunk_id)
                )
            ''')

            conn.commit()
            logger.info("Database initialized successfully")

    def insert_document(self, doc_data: Dict):
        """Insert document metadata."""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO documents 
                (document_id, filename, file_type, file_size_mb, page_count, 
                 word_count, character_count, upload_timestamp, status, error_message)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                doc_data['document_id'],
                doc_data['filename'],
                doc_data.get('file_type', ''),
                doc_data.get('file_size_mb', 0),
                doc_data.get('page_count', 0),
                doc_data.get('word_count', 0),
                doc_data.get('character_count', 0),
                doc_data.get('upload_timestamp', datetime.utcnow().isoformat()),
                doc_data.get('status', 'processing'),
                doc_data.get('error_message', None)
            ))

    def insert_chunk(self, chunk_data: Dict):
        """Insert chunk metadata."""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO chunks 
                (chunk_id, document_id, chunk_index, text, word_count, 
                 embedding_generated, stored_in_vectordb, created_timestamp)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                chunk_data['chunk_id'],
                chunk_data['document_id'],
                chunk_data['chunk_index'],
                chunk_data['text'],
                chunk_data.get('word_count', 0),
                chunk_data.get('embedding_generated', False),
                chunk_data.get('stored_in_vectordb', False),
                datetime.utcnow().isoformat()
            ))

    def insert_query(self, query_data: Dict):
        """Insert query metadata."""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO queries 
                (query_id, query_text, query_timestamp, num_results, 
                 llm_response, processing_time_ms)
                VALUES (?, ?, ?, ?, ?, ?)
            ''', (
                query_data['query_id'],
                query_data['query_text'],
                datetime.utcnow().isoformat(),
                query_data.get('num_results', 0),
                query_data.get('llm_response', ''),
                query_data.get('processing_time_ms', 0)
            ))

    def insert_query_result(self, query_id: str, chunk_id: str, score: float):
        """Insert query-chunk relationship."""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO query_results (query_id, chunk_id, relevance_score)
                VALUES (?, ?, ?)
            ''', (query_id, chunk_id, score))

    def get_all_documents(self):
        """Retrieve all documents."""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute('SELECT * FROM documents ORDER BY upload_timestamp DESC')
            return [dict(row) for row in cursor.fetchall()]

    def get_document_by_id(self, document_id: str):
        """Get specific document."""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute('SELECT * FROM documents WHERE document_id = ?', (document_id,))
            row = cursor.fetchone()
            return dict(row) if row else None

    def get_chunks_by_document(self, document_id: str):
        """Get all chunks for a document."""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute('''
                SELECT * FROM chunks 
                WHERE document_id = ? 
                ORDER BY chunk_index
            ''', (document_id,))
            return [dict(row) for row in cursor.fetchall()]

    def get_recent_queries(self, limit: int = 10):
        """Get recent queries."""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute('''
                SELECT * FROM queries 
                ORDER BY query_timestamp DESC 
                LIMIT ?
            ''', (limit,))
            return [dict(row) for row in cursor.fetchall()]

    def delete_document(self, document_id: str):
        """Delete document and its chunks."""
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute('DELETE FROM chunks WHERE document_id = ?', (document_id,))
            cursor.execute('DELETE FROM documents WHERE document_id = ?', (document_id,))


# Initialize Database
db_manager = DatabaseManager(DATABASE_PATH)


# ============================================================================
# PYDANTIC MODELS
# ============================================================================

class ChunkingConfig(BaseModel):
    chunk_size: int = DEFAULT_CHUNK_SIZE
    chunk_overlap: int = DEFAULT_CHUNK_OVERLAP
    strategy: str = "sliding_window"


class QueryRequest(BaseModel):
    query: str
    top_k: int = MAX_CONTEXT_CHUNKS
    use_llm: bool = True
    filter: Optional[Dict] = None


class LLMResponse(BaseModel):
    query: str
    answer: str
    sources: List[Dict]
    processing_time_ms: float


# ============================================================================
# DOCUMENT PROCESSING
# ============================================================================

class DocumentProcessor:
    """Handles document processing with validation and limits."""

    def __init__(self):
        self.supported_formats = SUPPORTED_FORMATS
        self.max_pages = MAX_PAGES_PER_DOC

    async def process_document(self, file_path: Path, original_filename: str) -> Dict:
        """Process a single document with validation."""
        file_ext = Path(original_filename).suffix.lower()

        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported format: {file_ext}")

        # Extract text based on file type
        if file_ext == '.pdf':
            raw_text, page_count = await self._process_pdf(file_path)
        elif file_ext == '.docx':
            raw_text, page_count = await self._process_docx(file_path)
        elif file_ext == '.txt':
            raw_text, page_count = await self._process_txt(file_path)
        elif file_ext in ['.html', '.htm']:
            raw_text, page_count = await self._process_html(file_path)
        elif file_ext == '.csv':
            raw_text, page_count = await self._process_csv(file_path)
        elif file_ext == '.json':
            raw_text, page_count = await self._process_json(file_path)
        elif file_ext == '.md':
            raw_text, page_count = await self._process_markdown(file_path)
        elif file_ext == '.xml':
            raw_text, page_count = await self._process_xml(file_path)
        else:
            raw_text, page_count = await self._process_txt(file_path)

        if page_count > self.max_pages:
            raise ValueError(f"Document exceeds {self.max_pages} pages (found {page_count})")

        cleaned_text = self._clean_text(raw_text)

        return {
            'text': cleaned_text,
            'page_count': page_count,
            'character_count': len(cleaned_text),
            'word_count': len(cleaned_text.split()),
        }

    async def _process_pdf(self, file_path: Path) -> tuple:
        text = ""
        page_count = 0
        try:
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        return text, page_count

    async def _process_docx(self, file_path: Path) -> tuple:
        doc = DocxDocument(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += " " + cell.text
            text += "\n"
        page_count = max(1, len(text.split()) // 500)
        return text, page_count

    async def _process_txt(self, file_path: Path) -> tuple:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as file:
                text = file.read()
        page_count = max(1, len(text.split()) // 500)
        return text, page_count

    async def _process_html(self, file_path: Path) -> tuple:
        with open(file_path, 'r', encoding='utf-8') as file:
            soup = BeautifulSoup(file.read(), 'html.parser')
            for script in soup(["script", "style"]):
                script.decompose()
            text = soup.get_text()
        page_count = max(1, len(text.split()) // 500)
        return text, page_count

    async def _process_csv(self, file_path: Path) -> tuple:
        df = pd.read_csv(file_path)
        text = df.to_string(index=False)
        page_count = max(1, len(text.split()) // 500)
        return text, page_count

    async def _process_json(self, file_path: Path) -> tuple:
        with open(file_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
            text = self._json_to_text(data)
        page_count = max(1, len(text.split()) // 500)
        return text, page_count

    async def _process_markdown(self, file_path: Path) -> tuple:
        text, _ = await self._process_txt(file_path)
        text = re.sub(r'#+ ', '', text)
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'`(.+?)`', r'\1', text)
        text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
        page_count = max(1, len(text.split()) // 500)
        return text, page_count

    async def _process_xml(self, file_path: Path) -> tuple:
        with open(file_path, 'r', encoding='utf-8') as file:
            soup = BeautifulSoup(file.read(), 'xml')
            text = soup.get_text()
        page_count = max(1, len(text.split()) // 500)
        return text, page_count

    def _json_to_text(self, data, prefix="") -> str:
        text = ""
        if isinstance(data, dict):
            for key, value in data.items():
                if isinstance(value, (dict, list)):
                    text += f"{prefix}{key}:\n"
                    text += self._json_to_text(value, prefix + "  ")
                else:
                    text += f"{prefix}{key}: {value}\n"
        elif isinstance(data, list):
            for item in data:
                text += self._json_to_text(item, prefix) + "\n"
        else:
            text += f"{prefix}{data}\n"
        return text

    def _clean_text(self, text: str) -> str:
        text = text.replace('\x00', '')
        text = re.sub(r'\r\n', '\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        text = re.sub(r'[^\x20-\x7E\n\t]', '', text)
        return text.strip()


# ============================================================================
# CHUNKING MODULE
# ============================================================================

class ChunkingModule:
    """Splits documents into optimally-sized chunks for retrieval."""

    def __init__(self):
        self.default_chunk_size = DEFAULT_CHUNK_SIZE
        self.default_overlap = DEFAULT_CHUNK_OVERLAP

    def chunk_documents(self, documents: List[Dict], config: ChunkingConfig) -> List[Dict]:
        """Chunk all documents based on strategy."""
        all_chunks = []

        for doc in documents:
            if doc['status'] != 'success':
                continue

            text = doc['text']
            doc_id = doc['document_id']
            filename = doc['filename']

            if config.strategy == "sliding_window":
                chunks = self._sliding_window_chunking(text, config.chunk_size, config.chunk_overlap)
            elif config.strategy == "sentence":
                chunks = self._sentence_chunking(text, config.chunk_size)
            elif config.strategy == "paragraph":
                chunks = self._paragraph_chunking(text, config.chunk_size)
            else:
                chunks = self._sliding_window_chunking(text, config.chunk_size, config.chunk_overlap)

            for idx, chunk_text in enumerate(chunks):
                all_chunks.append({
                    'chunk_id': f"{doc_id}_chunk_{idx}",
                    'document_id': doc_id,
                    'filename': filename,
                    'chunk_index': idx,
                    'text': chunk_text,
                    'char_count': len(chunk_text),
                    'word_count': len(chunk_text.split()),
                    'metadata': doc.get('metadata', {})
                })

        return all_chunks

    def _sliding_window_chunking(self, text: str, chunk_size: int, overlap: int) -> List[str]:
        """Split text using sliding window with overlap."""
        words = text.split()
        chunks = []

        for i in range(0, len(words), chunk_size - overlap):
            chunk = ' '.join(words[i:i + chunk_size])
            if chunk.strip():
                chunks.append(chunk)

            if i + chunk_size >= len(words):
                break

        return chunks

    def _sentence_chunking(self, text: str, max_chunk_size: int) -> List[str]:
        """Split by sentences, grouping to approximate chunk size."""
        sentences = re.split(r'(?<=[.!?])\s+', text)
        chunks = []
        current_chunk = []
        current_size = 0

        for sentence in sentences:
            sentence_words = len(sentence.split())

            if current_size + sentence_words > max_chunk_size and current_chunk:
                chunks.append(' '.join(current_chunk))
                current_chunk = [sentence]
                current_size = sentence_words
            else:
                current_chunk.append(sentence)
                current_size += sentence_words

        if current_chunk:
            chunks.append(' '.join(current_chunk))

        return chunks

    def _paragraph_chunking(self, text: str, max_chunk_size: int) -> List[str]:
        """Split by paragraphs, grouping to approximate chunk size."""
        paragraphs = text.split('\n\n')
        chunks = []
        current_chunk = []
        current_size = 0

        for para in paragraphs:
            para_words = len(para.split())

            if current_size + para_words > max_chunk_size and current_chunk:
                chunks.append('\n\n'.join(current_chunk))
                current_chunk = [para]
                current_size = para_words
            else:
                current_chunk.append(para)
                current_size += para_words

        if current_chunk:
            chunks.append('\n\n'.join(current_chunk))

        return chunks


# ============================================================================
# EMBEDDING MODULE
# ============================================================================

class EmbeddingModule:
    """Generates embeddings for text chunks."""

    def __init__(self, model_name: str = EMBEDDING_MODEL):
        logger.info(f"Loading embedding model: {model_name}")
        self.model = SentenceTransformer(model_name)
        self.embedding_dim = self.model.get_sentence_embedding_dimension()
        logger.info(f"Model loaded. Embedding dimension: {self.embedding_dim}")

    def generate_embeddings(self, chunks: List[Dict], batch_size: int = 32) -> List[Dict]:
        """Generate embeddings for all chunks."""
        texts = [chunk['text'] for chunk in chunks]

        logger.info(f"Generating embeddings for {len(texts)} chunks...")

        all_embeddings = []
        for i in range(0, len(texts), batch_size):
            batch = texts[i:i + batch_size]
            embeddings = self.model.encode(batch, show_progress_bar=False)
            all_embeddings.extend(embeddings)

        for chunk, embedding in zip(chunks, all_embeddings):
            chunk['embedding'] = embedding.tolist()
            chunk['embedding_model'] = EMBEDDING_MODEL
            chunk['embedding_dim'] = self.embedding_dim

        logger.info(f"Generated {len(all_embeddings)} embeddings")
        return chunks

    def embed_query(self, query: str) -> List[float]:
        """Generate embedding for a search query."""
        return self.model.encode(query).tolist()


# ============================================================================
# VECTOR STORE MODULE (PINECONE)
# ============================================================================

class PineconeVectorStore:
    """Manages Pinecone vector database operations."""

    def __init__(self, api_key: str, index_name: str):
        self.api_key = api_key
        self.index_name = index_name
        self.pc = None
        self.index = None
        self._initialize_pinecone()

    def _initialize_pinecone(self):
        """Initialize Pinecone connection and create/connect to index."""
        try:
            logger.info("Initializing Pinecone...")
            self.pc = Pinecone(api_key=self.api_key)

            existing_indexes = self.pc.list_indexes().names()

            if self.index_name not in existing_indexes:
                logger.info(f"Creating new index: {self.index_name}")
                self.pc.create_index(
                    name=self.index_name,
                    dimension=384,
                    metric="cosine",
                    spec=ServerlessSpec(
                        cloud="aws",
                        region="us-east-1"
                    )
                )
                logger.info("Index created successfully")
            else:
                logger.info(f"Connecting to existing index: {self.index_name}")

            self.index = self.pc.Index(self.index_name)
            logger.info("Pinecone initialized successfully")

        except Exception as e:
            logger.error(f"Pinecone initialization error: {e}")
            raise

    def upsert_chunks(self, chunks: List[Dict], batch_size: int = 100) -> Dict:
        """Upload chunks with embeddings to Pinecone."""
        vectors = []

        for chunk in chunks:
            vector_id = chunk['chunk_id']
            embedding = chunk['embedding']

            metadata = {
                'document_id': chunk['document_id'],
                'filename': chunk['filename'],
                'chunk_index': chunk['chunk_index'],
                'text': chunk['text'][:1000],
                'word_count': chunk['word_count'],
                'page_count': chunk['metadata'].get('page_count', 0),
            }

            vectors.append({
                'id': vector_id,
                'values': embedding,
                'metadata': metadata
            })

        logger.info(f"Upserting {len(vectors)} vectors to Pinecone...")
        upserted_count = 0

        for i in range(0, len(vectors), batch_size):
            batch = vectors[i:i + batch_size]
            self.index.upsert(vectors=batch)
            upserted_count += len(batch)
            logger.info(f"Upserted {upserted_count}/{len(vectors)} vectors")

        stats = self.index.describe_index_stats()

        return {
            'upserted_count': len(vectors),
            'total_vectors_in_index': stats.total_vector_count,
            'index_name': self.index_name
        }

    def search(self, query_embedding: List[float], top_k: int = 5, filter: Optional[Dict] = None) -> List[Dict]:
        """Search for similar chunks."""
        results = self.index.query(
            vector=query_embedding,
            top_k=top_k,
            include_metadata=True,
            filter=filter
        )

        return [
            {
                'chunk_id': match['id'],
                'score': match['score'],
                'metadata': match['metadata']
            }
            for match in results['matches']
        ]

    def delete_by_document_id(self, document_id: str):
        """Delete all chunks for a specific document."""
        self.index.delete(filter={'document_id': document_id})
        logger.info(f"Deleted vectors for document: {document_id}")


# ============================================================================
# LLM MODULE - GOOGLE GEMINI
# ============================================================================

class GeminiLLMModule:
    """Handles Google Gemini API calls for generating responses."""

    def __init__(self, api_key: str, model: str = GEMINI_MODEL):
        self.api_key = api_key
        self.model_name = model

        # Configure Gemini
        genai.configure(api_key=api_key)

        # Initialize the model
        self.model = genai.GenerativeModel(
            model_name=model,
            generation_config={
                "temperature": 0.3,
                "top_p": 0.95,
                "top_k": 40,
                "max_output_tokens": 2048,
            },
            safety_settings=[
                {
                    "category": "HARM_CATEGORY_HARASSMENT",
                    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
                },
                {
                    "category": "HARM_CATEGORY_HATE_SPEECH",
                    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
                },
                {
                    "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
                    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
                },
                {
                    "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
                    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
                },
            ]
        )

        logger.info(f"Gemini LLM Module initialized with model: {model}")

    def generate_response(self, query: str, context_chunks: List[Dict]) -> str:
        """Generate response using Google Gemini with retrieved context."""

        # Build context from retrieved chunks
        context = "\n\n".join([
            f"[Source: {chunk['metadata']['filename']}]\n{chunk['metadata']['text']}"
            for chunk in context_chunks
        ])

        # Create prompt
        prompt = f"""You are a helpful AI assistant. Answer the user's question based ONLY on the provided context from documents.

Rules:
- Be accurate, concise, and relevant
- If the answer is not in the context, say "I don't have enough information to answer that question."
- Cite the source document when possible
- Don't make up information

Context from documents:
{context}

Question: {query}

Answer:"""

        try:
            logger.info(f"Sending query to Gemini: {query[:100]}...")

            # Generate response
            response = self.model.generate_content(prompt)

            # Extract text from response
            if response.text:
                answer = response.text
                logger.info("Gemini response received successfully")
            else:
                answer = "I apologize, but I couldn't generate a response. Please try rephrasing your question."
                logger.warning("Gemini returned empty response")

            return answer

        except Exception as e:
            logger.error(f"Gemini API error: {e}")
            return f"Error generating response: {str(e)}"


# ============================================================================
# COMPLETE RAG PIPELINE
# ============================================================================

class CompleteRAGPipeline:
    """Orchestrates the complete RAG system."""

    def __init__(self):
        self.processor = DocumentProcessor()
        self.chunker = ChunkingModule()
        self.embedder = EmbeddingModule()
        self.vectorstore = PineconeVectorStore(PINECONE_API_KEY, PINECONE_INDEX_NAME)
        self.llm = GeminiLLMModule(GEMINI_API_KEY, GEMINI_MODEL)
        self.db = db_manager

        self.pipeline_state = {}

    async def upload_and_process(self, files: List[UploadFile], chunking_config: ChunkingConfig) -> Dict:
        """Complete pipeline: Upload → Process → Chunk → Embed → Store."""

        results = []
        temp_dir = Path(tempfile.mkdtemp())

        try:
            # Step 1: Process documents
            for idx, file in enumerate(files):
                logger.info(f"Processing file {idx + 1}/{len(files)}: {file.filename}")

                temp_file = temp_dir / f"{uuid.uuid4()}_{file.filename}"
                with open(temp_file, 'wb') as f:
                    content = await file.read()
                    f.write(content)

                file_size_mb = temp_file.stat().st_size / (1024 * 1024)
                if file_size_mb > MAX_FILE_SIZE_MB:
                    raise ValueError(f"{file.filename}: Exceeds {MAX_FILE_SIZE_MB}MB limit")

                document_id = str(uuid.uuid4())

                try:
                    # Process document
                    processed = await self.processor.process_document(temp_file, file.filename)

                    doc_data = {
                        'document_id': document_id,
                        'filename': file.filename,
                        'file_type': Path(file.filename).suffix.lower(),
                        'file_size_mb': round(file_size_mb, 2),
                        'page_count': processed['page_count'],
                        'word_count': processed['word_count'],
                        'character_count': processed['character_count'],
                        'upload_timestamp': datetime.utcnow().isoformat(),
                        'status': 'success',
                        'error_message': None
                    }

                    # Save to database
                    self.db.insert_document(doc_data)

                    results.append({
                        **doc_data,
                        'text': processed['text'],
                        'metadata': {
                            'page_count': processed['page_count'],
                            'word_count': processed['word_count'],
                            'character_count': processed['character_count']
                        }
                    })

                except Exception as e:
                    logger.error(f"Error processing {file.filename}: {e}")

                    # Save error to database
                    self.db.insert_document({
                        'document_id': document_id,
                        'filename': file.filename,
                        'file_type': Path(file.filename).suffix.lower(),
                        'file_size_mb': round(file_size_mb, 2),
                        'status': 'failed',
                        'error_message': str(e)
                    })

                    results.append({
                        'document_id': document_id,
                        'filename': file.filename,
                        'status': 'failed',
                        'error': str(e)
                    })

            successful_docs = [r for r in results if r['status'] == 'success']

            if not successful_docs:
                return {
                    'status': 'failed',
                    'message': 'No documents processed successfully',
                    'results': results
                }

            # Step 2: Chunk documents
            chunks = self.chunker.chunk_documents(successful_docs, chunking_config)

            # Save chunks to database
            for chunk in chunks:
                self.db.insert_chunk({
                    'chunk_id': chunk['chunk_id'],
                    'document_id': chunk['document_id'],
                    'chunk_index': chunk['chunk_index'],
                    'text': chunk['text'],
                    'word_count': chunk['word_count'],
                    'embedding_generated': False,
                    'stored_in_vectordb': False
                })

            # Step 3: Generate embeddings
            embedded_chunks = self.embedder.generate_embeddings(chunks)

            # Step 4: Store in Pinecone
            vectorstore_result = self.vectorstore.upsert_chunks(embedded_chunks)

            return {
                'status': 'success',
                'message': 'Documents processed and stored successfully',
                'summary': {
                    'documents_uploaded': len(files),
                    'documents_processed': len(successful_docs),
                    'documents_failed': len([r for r in results if r['status'] == 'failed']),
                    'total_chunks': len(chunks),
                    'embeddings_generated': len(embedded_chunks),
                    'vectors_stored': vectorstore_result['upserted_count']
                },
                'documents': results
            }

        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

    async def query_with_llm(self, query: str, top_k: int = MAX_CONTEXT_CHUNKS,
                             use_llm: bool = True, filter: Optional[Dict] = None) -> Dict:
        """Query the system and optionally generate LLM response."""

        start_time = datetime.utcnow()
        query_id = str(uuid.uuid4())

        # Generate query embedding
        query_embedding = self.embedder.embed_query(query)

        # Search in Pinecone
        search_results = self.vectorstore.search(query_embedding, top_k, filter)

        # Prepare response
        response_data = {
            'query_id': query_id,
            'query': query,
            'num_results': len(search_results),
            'results': search_results
        }

        # Generate LLM response if requested
        if use_llm and search_results:
            llm_answer = self.llm.generate_response(query, search_results)
            response_data['llm_response'] = llm_answer
        else:
            response_data['llm_response'] = None

        # Calculate processing time
        processing_time_ms = (datetime.utcnow() - start_time).total_seconds() * 1000
        response_data['processing_time_ms'] = processing_time_ms

        # Save query to database
        self.db.insert_query({
            'query_id': query_id,
            'query_text': query,
            'num_results': len(search_results),
            'llm_response': response_data.get('llm_response', ''),
            'processing_time_ms': processing_time_ms
        })

        # Save query-chunk relationships
        for result in search_results:
            self.db.insert_query_result(query_id, result['chunk_id'], result['score'])

        return response_data


# Initialize pipeline
pipeline = CompleteRAGPipeline()


# ============================================================================
# API ENDPOINTS
# ============================================================================

@app.get("/")
async def root():
    return {
        "status": "operational",
        "service": "Complete RAG System with Gemini LLM",
        "version": "3.1.0",
        "features": [
            "Document Upload & Processing",
            "Semantic Search with Pinecone",
            "AI-Powered Responses with Google Gemini",
            "Metadata Storage in SQLite"
        ],
        "endpoints": {
            "upload": "/api/v1/documents/upload",
            "query": "/api/v1/query",
            "documents": "/api/v1/documents",
            "health": "/api/v1/health",
            "docs": "/docs"
        }
    }


@app.post("/api/v1/documents/upload")
async def upload_documents(
        files: List[UploadFile] = File(...),
        chunk_size: int = DEFAULT_CHUNK_SIZE,
        chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
        chunking_strategy: str = "sliding_window"
):
    """
    Upload and process documents.

    This endpoint:
    1. Accepts document files
    2. Extracts and cleans text
    3. Chunks the text
    4. Generates embeddings
    5. Stores in Pinecone vector database
    6. Saves metadata to SQLite database
    """

    if len(files) > MAX_DOCUMENTS:
        raise HTTPException(status_code=400, detail=f"Max {MAX_DOCUMENTS} documents allowed")

    if len(files) == 0:
        raise HTTPException(status_code=400, detail="No files provided")

    for file in files:
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in SUPPORTED_FORMATS:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported format '{file_ext}' in '{file.filename}'"
            )

    try:
        chunking_config = ChunkingConfig(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            strategy=chunking_strategy
        )

        result = await pipeline.upload_and_process(files, chunking_config)
        return JSONResponse(content=result, status_code=200)

    except Exception as e:
        logger.error(f"Upload error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/v1/query")
async def query_system(request: QueryRequest):
    """
    Query the RAG system.

    This endpoint:
    1. Accepts a user query
    2. Converts query to embedding
    3. Searches for relevant document chunks
    4. Optionally generates AI response using Google Gemini
    5. Returns results with sources

    Example request:
    {
        "query": "What is machine learning?",
        "top_k": 5,
        "use_llm": true,
        "filter": {"filename": "ml_guide.pdf"}
    }
    """

    try:
        result = await pipeline.query_with_llm(
            query=request.query,
            top_k=request.top_k,
            use_llm=request.use_llm,
            filter=request.filter
        )
        return JSONResponse(content=result, status_code=200)

    except Exception as e:
        logger.error(f"Query error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/v1/documents")
async def get_all_documents():
    """
    Get metadata for all uploaded documents.

    Returns:
    - List of all documents with their metadata
    - Processing status
    - Statistics (page count, word count, etc.)
    """

    try:
        documents = pipeline.db.get_all_documents()

        return {
            'total_documents': len(documents),
            'successful': len([d for d in documents if d['status'] == 'success']),
            'failed': len([d for d in documents if d['status'] == 'failed']),
            'documents': documents
        }

    except Exception as e:
        logger.error(f"Error fetching documents: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/v1/documents/{document_id}")
async def get_document_details(document_id: str):
    """
    Get detailed information about a specific document.

    Returns:
    - Document metadata
    - All chunks for the document
    """

    try:
        document = pipeline.db.get_document_by_id(document_id)

        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        chunks = pipeline.db.get_chunks_by_document(document_id)

        return {
            'document': document,
            'chunks': chunks,
            'total_chunks': len(chunks)
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching document: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/v1/documents/{document_id}")
async def delete_document(document_id: str):
    """
    Delete a document and all its chunks.

    This removes:
    - Document metadata from database
    - All chunk records
    - Vectors from Pinecone
    """

    try:
        # Check if document exists
        document = pipeline.db.get_document_by_id(document_id)
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        # Delete from vector store
        pipeline.vectorstore.delete_by_document_id(document_id)

        # Delete from database
        pipeline.db.delete_document(document_id)

        return {
            'status': 'success',
            'message': f'Document {document_id} deleted successfully',
            'filename': document['filename']
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Delete error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/v1/queries/recent")
async def get_recent_queries(limit: int = 10):
    """
    Get recent queries and their responses.

    Useful for:
    - Monitoring system usage
    - Analyzing common questions
    - Debugging
    """

    try:
        queries = pipeline.db.get_recent_queries(limit)

        return {
            'total_queries': len(queries),
            'queries': queries
        }

    except Exception as e:
        logger.error(f"Error fetching queries: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/v1/health")
async def health_check():
    """
    Comprehensive health check.

    Checks:
    - API status
    - Pinecone connection
    - Database connection
    - Google Gemini API availability
    """

    health_status = {
        "status": "healthy",
        "timestamp": datetime.utcnow().isoformat(),
        "version": "3.1.0"
    }

    # Check Pinecone
    try:
        index_stats = pipeline.vectorstore.index.describe_index_stats()
        health_status['pinecone'] = {
            'status': 'connected',
            'index_name': PINECONE_INDEX_NAME,
            'total_vectors': index_stats.total_vector_count
        }
    except Exception as e:
        health_status['pinecone'] = {
            'status': 'error',
            'error': str(e)
        }

    # Check Database
    try:
        documents = pipeline.db.get_all_documents()
        health_status['database'] = {
            'status': 'connected',
            'total_documents': len(documents),
            'database_path': DATABASE_PATH
        }
    except Exception as e:
        health_status['database'] = {
            'status': 'error',
            'error': str(e)
        }

    # Check Gemini (simple check)
    try:
        health_status['gemini'] = {
            'status': 'configured',
            'model': GEMINI_MODEL
        }
    except Exception as e:
        health_status['gemini'] = {
            'status': 'error',
            'error': str(e)
        }

    # Configuration
    health_status['configuration'] = {
        'max_documents': MAX_DOCUMENTS,
        'max_pages_per_doc': MAX_PAGES_PER_DOC,
        'max_file_size_mb': MAX_FILE_SIZE_MB,
        'supported_formats': list(SUPPORTED_FORMATS),
        'chunk_size': DEFAULT_CHUNK_SIZE,
        'chunk_overlap': DEFAULT_CHUNK_OVERLAP,
        'embedding_model': EMBEDDING_MODEL
    }

    return health_status


@app.get("/api/v1/stats")
async def get_statistics():
    """
    Get overall system statistics.

    Provides insights on:
    - Total documents processed
    - Total chunks stored
    - Recent activity
    """

    try:
        # Get documents
        documents = pipeline.db.get_all_documents()

        # Get Pinecone stats
        index_stats = pipeline.vectorstore.index.describe_index_stats()

        # Get recent queries
        recent_queries = pipeline.db.get_recent_queries(5)

        return {
            'timestamp': datetime.utcnow().isoformat(),
            'documents': {
                'total': len(documents),
                'successful': len([d for d in documents if d['status'] == 'success']),
                'failed': len([d for d in documents if d['status'] == 'failed'])
            },
            'vectorstore': {
                'total_vectors': index_stats.total_vector_count,
                'index_name': PINECONE_INDEX_NAME,
                'dimension': 384
            },
            'recent_queries': len(recent_queries),
            'models': {
                'embedding': EMBEDDING_MODEL,
                'llm': GEMINI_MODEL
            }
        }

    except Exception as e:
        logger.error(f"Stats error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# RUN SERVER
# ============================================================================

if __name__ == "__main__":
    import uvicorn
    import nest_asyncio

    print("=" * 80)
    print("🚀 COMPLETE RAG SYSTEM WITH GOOGLE GEMINI LLM")
    print("=" * 80)
    print("\n📊 CONFIGURATION:")
    print(f"   Supported Formats: {', '.join(SUPPORTED_FORMATS)}")
    print(f"   Max Documents: {MAX_DOCUMENTS}")
    print(f"   Max File Size: {MAX_FILE_SIZE_MB}MB")
    print(f"   Embedding Model: {EMBEDDING_MODEL}")
    print(f"   LLM Model: {GEMINI_MODEL}")
    print(f"   Vector Database: Pinecone ({PINECONE_INDEX_NAME})")
    print(f"   Metadata Database: SQLite ({DATABASE_PATH})")

    print("\n🔄 PIPELINE STEPS:")
    print("   1. Document Upload & Text Extraction")
    print("   2. Text Chunking (Configurable)")
    print("   3. Embedding Generation (SentenceTransformers)")
    print("   4. Vector Storage (Pinecone)")
    print("   5. Metadata Storage (SQLite)")
    print("   6. Semantic Search")
    print("   7. LLM Response Generation (Google Gemini)")

    print("\n" + "=" * 80)
    print("📡 SERVER STARTING...")
    print("=" * 80)
    print("   🌐 Base URL:    http://localhost:8000")
    print("   📚 API Docs:    http://localhost:8000/docs")
    print("   ❤️  Health:      http://localhost:8000/api/v1/health")
    print("   📊 Stats:       http://localhost:8000/api/v1/stats")

    print("\n" + "=" * 80)
    print("🎯 KEY ENDPOINTS:")
    print("=" * 80)
    print("   📤 Upload:      POST /api/v1/documents/upload")
    print("   🔍 Query:       POST /api/v1/query")
    print("   📋 Documents:   GET  /api/v1/documents")
    print("   🗑️  Delete:      DELETE /api/v1/documents/{id}")
    print("   📜 Queries:     GET  /api/v1/queries/recent")
    print("=" * 80)
    print("\n✨ System ready! Open http://localhost:8000/docs to start.\n")

    nest_asyncio.apply()
    uvicorn.run(app, host="0.0.0.0", port=8000)